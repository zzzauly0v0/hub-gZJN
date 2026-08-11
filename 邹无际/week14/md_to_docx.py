# -*- coding: utf-8 -*-
"""MD → DOCX 转换器（统一格式）

用法:
    python md_to_docx.py <md文件...> [--out <输出docx路径或目录>]

说明:
    - 每个 md 生成同名 docx（与 md 同目录，文档标题 = md 文件名）
    - 输出文件被占用时自动追加 _v2 重试
    - 支持一次转换多个文件
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

F, FC = '微软雅黑', 'Consolas'
SZ = dict(t=Pt(22), h2=Pt(16), h3=Pt(14), b=Pt(12), th=Pt(11.5), td=Pt(11), q=Pt(12))
CL = dict(t=RGBColor(0x1A, 0x3C, 0x6E), h2=RGBColor(0x1A, 0x3C, 0x6E),
          h3=RGBColor(0x2B, 0x57, 0x97), code=RGBColor(0xC7, 0x25, 0x4E))


def style(r, font=F, size=SZ['b'], bold=False, color=None):
    r.font.name, r.font.size, r.bold = font, size, bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        r.font.color.rgb = color


def add(p, text, **kw):
    style(p.add_run(text), **kw)


def shade(cell, hexc):
    tc = cell._element.get_or_add_tcPr()
    for o in tc.findall(qn('w:shd')):
        tc.remove(o)
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): hexc, qn('w:val'): 'clear'}))


def inline(p, text, size=SZ['b'], bold=False, color=None):
    last = 0
    for m in re.finditer(r'\*\*(.+?)\*\*|`(.+?)`', text):
        if m.start() > last:
            add(p, text[last:m.start()], size=size, bold=bold, color=color)
        if m.group(1):
            add(p, m.group(1), size=size, bold=True, color=color)
        else:
            add(p, m.group(2), font=FC, size=Pt(size.pt - 0.5), color=CL['code'])
        last = m.end()
    if last < len(text):
        add(p, text[last:], size=size, bold=bold, color=color)


def para(doc, sb=4, sa=4, indent=None, style_name=None):
    p = doc.add_paragraph(style=style_name)
    pf = p.paragraph_format
    pf.space_before, pf.space_after, pf.line_spacing_rule = Pt(sb), Pt(sa), WD_LINE_SPACING.SINGLE
    if indent:
        pf.left_indent = indent
    return p


def heading(doc, text, level, size, color, sb, sa):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before, h.paragraph_format.space_after = Pt(sb), Pt(sa)
    for r in h.runs:
        style(r, size=size, bold=True, color=color)
    return h


def add_table(doc, hdrs, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(hdrs))
    t.style, t.alignment = 'Table Grid', WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(hdrs):
        c = t.rows[0].cells[ci]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(5), Pt(5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        add(p, h, size=SZ['th'], bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(c, '1A3C6E')
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(4), Pt(4)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            inline(p, v, size=SZ['td'])
            if ri % 2:
                shade(c, 'E8EFF8')
    para(doc, 2, 2)


def add_quote(doc, text):
    p = para(doc, 3, 3, indent=Cm(1.0))
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    bdr = pPr.makeelement(qn('w:pBdr'), {})
    bdr.append(bdr.makeelement(qn('w:left'), {qn('w:val'): 'single', qn('w:sz'): '24',
                                              qn('w:space'): '8', qn('w:color'): '2B5797'}))
    pPr.append(bdr)
    pPr.append(pPr.makeelement(qn('w:shd'), {qn('w:fill'): 'EDF2F9', qn('w:val'): 'clear'}))
    inline(p, text, size=SZ['q'])


def parse_table(buf):
    hdrs = [c.strip() for c in buf[0].strip('|').split('|')]
    rows = [[c.strip() for c in l.strip('|').split('|')] for l in buf[2:]]
    return hdrs, rows


def convert(md, docx):
    with open(md, encoding='utf-8') as f:
        lines = f.read().split('\n')
    doc = Document()
    ns = doc.styles['Normal']
    style(ns, size=SZ['b'])
    ns.paragraph_format.space_before, ns.paragraph_format.space_after = Pt(4), Pt(4)
    ns.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for lvl in range(1, 5):
        style(doc.styles[f'Heading {lvl}'], bold=True)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.5)
    tp = doc.add_heading(os.path.splitext(os.path.basename(md))[0], level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before, tp.paragraph_format.space_after = Pt(24), Pt(16)
    for r in tp.runs:
        style(r, size=SZ['t'], bold=True, color=CL['t'])
    buf, skipped = [], False
    for line in lines[1:]:
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            buf.append(line)
            continue
        if buf:
            hdrs, rows = parse_table(buf)
            add_table(doc, hdrs, rows)
            buf = []
        if line.strip() == '---':
            p = para(doc, 10, 10)
            bdr = p._element.get_or_add_pPr().makeelement(qn('w:pBdr'), {})
            bdr.append(bdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '12',
                                                        qn('w:space'): '1', qn('w:color'): '999999'}))
            p._element.get_or_add_pPr().append(bdr)
        elif line.startswith('# ') and not skipped:
            skipped = True
        elif line.startswith('## '):
            heading(doc, line[3:].strip(), 2, SZ['h2'], CL['h2'], 20, 10)
        elif line.startswith('### '):
            heading(doc, line[4:].strip(), 3, SZ['h3'], CL['h3'], 16, 8)
        elif line.startswith('>'):
            q = line.lstrip('> ').strip()
            if q:
                add_quote(doc, q)
            else:
                para(doc, 1, 1, indent=Cm(1.0))
        elif line.strip().startswith('- '):
            p = para(doc, 2, 2, style_name='List Bullet')
            inline(p, line.strip()[2:])
        elif re.match(r'^\s*\d+\.\s+', line):
            p = para(doc, 2, 2, style_name='List Number')
            inline(p, re.sub(r'^\s*\d+\.\s+', '', line))
        elif line.strip():
            p = para(doc)
            inline(p, line.strip())
    if buf:
        hdrs, rows = parse_table(buf)
        add_table(doc, hdrs, rows)
    try:
        doc.save(docx)
    except PermissionError:
        base, ext = os.path.splitext(docx)
        docx = f'{base}_v2{ext}'
        doc.save(docx)
    print(f'已生成: {docx}')


def main():
    argv = sys.argv[1:]
    out = None
    if '--out' in argv:
        i = argv.index('--out')
        out, argv = argv[i + 1], argv[:i] + argv[i + 2:]
    if not argv:
        print(__doc__)
        return
    for md in argv:
        if not os.path.isfile(md):
            print(f'文件不存在，跳过: {md}')
            continue
        docx = out
        if not docx:
            docx = os.path.splitext(md)[0] + '.docx'
        elif os.path.isdir(docx) or not os.path.splitext(docx)[1]:
            os.makedirs(docx, exist_ok=True)
            docx = os.path.join(docx, os.path.splitext(os.path.basename(md))[0] + '.docx')
        convert(md, docx)


if __name__ == '__main__':
    main()
